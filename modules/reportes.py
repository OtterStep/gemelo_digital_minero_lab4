"""
Módulo de Reportes Multi-formato
Generación de reportes en PDF, Word y Excel.
"""
import pandas as pd
import numpy as np
import io
import os
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.utils.dataframe import dataframe_to_rows
import plotly.io as pio

from utils.database import get_connection
from modules import dashboard as dash_mod

# ============================================================
# GENERACIÓN DE REPORTES PDF
# ============================================================

def generar_reporte_pdf(tipo_reporte='ejecutivo', periodo='Últimos 30 días'):
    """Generar reporte ejecutivo en PDF"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           rightMargin=0.75*inch, leftMargin=0.75*inch,
                           topMargin=1*inch, bottomMargin=0.75*inch)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Estilos personalizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=20,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=20,
        alignment=TA_CENTER
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#34495e'),
        spaceAfter=15,
        spaceBefore=10
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=8
    )
    
    # Título
    story.append(Paragraph("SISTEMA DE GEMELOS DIGITALES", title_style))
    story.append(Paragraph("Reporte Ejecutivo de Mantenimiento de Equipos de Carguío", subtitle_style))
    story.append(Paragraph(f"Periodo: {periodo} | Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}", normal_style))
    story.append(Spacer(1, 0.2*inch))
    
    # KPIs
    story.append(Paragraph("INDICADORES CLAVE DE DESEMPEÑO (KPIs)", subtitle_style))
    
    kpis = dash_mod.get_kpis()
    
    kpi_data = [
        ['Indicador', 'Valor', 'Meta'],
        ['Disponibilidad de Equipos', f"{kpis['disponibilidad']}%", '≥ 90%'],
        ['OEE (Efectividad General)', f"{kpis['oee']}%", '≥ 75%'],
        ['MTBF (Horas entre fallas)', f"{kpis['mtbf']:.0f} h", '≥ 500 h'],
        ['MTTR (Tiempo de reparación)', f"{kpis['mttr']:.1f} h", '≤ 8 h'],
        ['Costos de Mantenimiento', f"${kpis['costos_mantenimiento']:,.2f}", 'Presupuestado'],
        ['Equipos Críticos', str(kpis['equipos_criticos']), '0'],
        ['Órdenes Pendientes', str(kpis['ordenes_pendientes']), '≤ 5'],
    ]
    
    kpi_table = Table(kpi_data, colWidths=[3*inch, 1.5*inch, 1.5*inch])
    kpi_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ecf0f1')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
    ]))
    
    story.append(kpi_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Análisis de Equipos
    story.append(Paragraph("ANÁLISIS DE FLOTA DE EQUIPOS", subtitle_style))
    
    equipos_df = dash_mod.get_equipos_data()
    estado_data = [['Estado', 'Cantidad']]
    for estado, count in equipos_df['estado'].value_counts().items():
        estado_data.append([estado, count])
    
    estado_table = Table(estado_data, colWidths=[2.5*inch, 2.5*inch])
    estado_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
    ]))
    
    story.append(estado_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Órdenes de Trabajo Recientes
    story.append(Paragraph("ÓRDENES DE TRABAJO RECIENTES", subtitle_style))
    
    ordenes_df = dash_mod.get_ordenes_trabajo_data().head(10)
    
    if not ordenes_df.empty:
        ot_data = [['N° Orden', 'Equipo', 'Tipo', 'Prioridad', 'Estado', 'Costo Real']]
        for _, row in ordenes_df.iterrows():
            ot_data.append([
                row['numero_orden'],
                row['equipo_codigo'],
                row['tipo'],
                row['prioridad'],
                row['estado'],
                f"${row['costo_real']:,.2f}"
            ])
        
        ot_table = Table(ot_data, colWidths=[1.2*inch, 1*inch, 0.9*inch, 0.9*inch, 1*inch, 1*inch])
        ot_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e74c3c')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
        ]))
        story.append(ot_table)
    
    story.append(Spacer(1, 0.3*inch))
    
    # Recomendaciones
    story.append(Paragraph("RECOMENDACIONES", subtitle_style))
    
    recomendaciones = [
        "• Los equipos con MTBF por debajo de la meta requieren análisis de causa raíz inmediato.",
        "• Implementar programa de mantenimiento predictivo para reducir MTTR.",
        "• Optimizar inventario de repuestos críticos para minimizar tiempos de espera.",
        "• Realizar capacitación continua a técnicos en nuevas tecnologías de diagnóstico.",
        "• Monitorear tendencias de consumo de combustible para detectar anomalías tempranas."
    ]
    
    for rec in recomendaciones:
        story.append(Paragraph(rec, normal_style))
    
    story.append(Spacer(1, 0.5*inch))
    
    # Pie de página
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#7f8c8d'),
        alignment=TA_CENTER
    )
    story.append(Paragraph("-" * 80, footer_style))
    story.append(Paragraph("Sistema de Gemelos Digitales para Minería | Generado automáticamente", footer_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

# ============================================================
# GENERACIÓN DE REPORTES WORD
# ============================================================

def generar_reporte_word(equipo_id=None):
    """Generar reporte detallado en Word por equipo o general"""
    document = Document()
    
    # Configurar estilos
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Título
    title = document.add_heading('INFORME DETALLADO DE MANTENIMIENTO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitle = document.add_heading('Sistema de Gemelos Digitales - Equipos de Carguío', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información general
    document.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    document.add_paragraph(f'Elaborado por: Sistema Automatizado de Reportes')
    
    document.add_heading('1. RESUMEN EJECUTIVO', level=1)
    
    kpis = dash_mod.get_kpis()
    
    resumen_text = f"""
    El presente informe detalla el estado actual de la flota de equipos de carguío, 
    indicadores de mantenimiento y recomendaciones para la optimización de operaciones.
    
    • Disponibilidad de flota: {kpis['disponibilidad']}%
    • Efectividad General de Equipos (OEE): {kpis['oee']}%
    • Tiempo Medio Entre Fallas (MTBF): {kpis['mtbf']:.0f} horas
    • Tiempo Medio de Reparación (MTTR): {kpis['mttr']:.1f} horas
    • Costos totales de mantenimiento: ${kpis['costos_mantenimiento']:,.2f}
    """
    document.add_paragraph(resumen_text)
    
    document.add_heading('2. ESTADO DE LA FLOTA', level=1)
    
    equipos_df = dash_mod.get_equipos_data()
    
    # Tabla de equipos
    document.add_heading('2.1. Distribución por Estado', level=2)
    
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    headers = ['Código', 'Nombre', 'Tipo', 'Horas Operación', 'Estado']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for _, eq in equipos_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = eq['codigo']
        row_cells[1].text = eq['nombre']
        row_cells[2].text = eq['tipo']
        row_cells[3].text = f"{eq['horas_operacion']:.0f}"
        row_cells[4].text = eq['estado']
    
    document.add_heading('3. ÓRDENES DE TRABAJO', level=1)
    
    ordenes_df = dash_mod.get_ordenes_trabajo_data()
    
    document.add_heading('3.1. Resumen por Tipo', level=2)
    for tipo, count in ordenes_df['tipo'].value_counts().items():
        document.add_paragraph(f'• {tipo}: {count} órdenes', style='List Bullet')
    
    document.add_heading('3.2. Detalle de Órdenes', level=2)
    
    if not ordenes_df.empty:
        table2 = document.add_table(rows=1, cols=6)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        table2.style = 'Light Grid Accent 1'
        
        hdr2 = table2.rows[0].cells
        headers2 = ['N° Orden', 'Equipo', 'Tipo', 'Prioridad', 'Estado', 'Costo Real']
        for i, h in enumerate(headers2):
            hdr2[i].text = h
            hdr2[i].paragraphs[0].runs[0].font.bold = True
        
        for _, ot in ordenes_df.head(15).iterrows():
            row = table2.add_row().cells
            row[0].text = ot['numero_orden']
            row[1].text = ot['equipo_codigo']
            row[2].text = ot['tipo']
            row[3].text = ot['prioridad']
            row[4].text = ot['estado']
            row[5].text = f"${ot['costo_real']:,.2f}"
    
    document.add_heading('4. ANÁLISIS DE COSTOS', level=1)
    
    costo_total = ordenes_df['costo_real'].sum()
    costo_estimado = ordenes_df['costo_estimado'].sum()
    
    document.add_paragraph(f'• Costo total real de mantenimiento: ${costo_total:,.2f}')
    document.add_paragraph(f'• Costo estimado total: ${costo_estimado:,.2f}')
    if costo_estimado > 0:
        variacion = ((costo_total - costo_estimado) / costo_estimado) * 100
        document.add_paragraph(f'• Variación: {variacion:+.1f}%')
    
    document.add_heading('5. CONCLUSIONES Y RECOMENDACIONES', level=1)
    
    conclusiones = [
        'La disponibilidad de la flota se encuentra dentro de los parámetros aceptables.',
        'Se identifican oportunidades de mejora en el tiempo medio de reparación (MTTR).',
        'El mantenimiento preventivo ha contribuido a reducir las fallas no programadas.',
        'Se recomienda implementar sensores adicionales para monitoreo de vibraciones.'
    ]
    
    for c in conclusiones:
        document.add_paragraph(c, style='List Number')
    
    # Guardar en buffer
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ============================================================
# GENERACIÓN DE REPORTES EXCEL
# ============================================================

def generar_reporte_excel():
    """Generar reporte completo en Excel con múltiples hojas"""
    wb = Workbook()
    
    # Estilos
    header_font = Font(name='Calibri', size=11, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='2C3E50', end_color='2C3E50', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    def apply_header_style(ws, max_col):
        for col in range(1, max_col + 1):
            cell = ws.cell(row=1, column=col)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
    
    def auto_width(ws):
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    
    # ============================================================
    # Hoja 1: KPIs
    # ============================================================
    ws1 = wb.active
    ws1.title = 'KPIs'
    
    kpis = dash_mod.get_kpis()
    kpi_data = [
        ['Indicador', 'Valor', 'Meta', 'Estado'],
        ['Disponibilidad de Equipos (%)', kpis['disponibilidad'], 90, '✓' if kpis['disponibilidad'] >= 90 else '⚠'],
        ['OEE (%)', kpis['oee'], 75, '✓' if kpis['oee'] >= 75 else '⚠'],
        ['MTBF (horas)', kpis['mtbf'], 500, '✓' if kpis['mtbf'] >= 500 else '⚠'],
        ['MTTR (horas)', kpis['mttr'], 8, '✓' if kpis['mttr'] <= 8 else '⚠'],
        ['Costos Mantenimiento ($)', kpis['costos_mantenimiento'], 'Presupuesto', '-'],
        ['Equipos Críticos', kpis['equipos_criticos'], 0, '✓' if kpis['equipos_criticos'] == 0 else '⚠'],
        ['Órdenes Pendientes', kpis['ordenes_pendientes'], 5, '✓' if kpis['ordenes_pendientes'] <= 5 else '⚠'],
    ]
    
    for row in kpi_data:
        ws1.append(row)
    
    apply_header_style(ws1, 4)
    auto_width(ws1)
    
    # ============================================================
    # Hoja 2: Equipos
    # ============================================================
    ws2 = wb.create_sheet('Equipos')
    
    equipos_df = dash_mod.get_equipos_data()
    for r in dataframe_to_rows(equipos_df, index=False, header=True):
        ws2.append(r)
    
    apply_header_style(ws2, len(equipos_df.columns))
    auto_width(ws2)
    
    # ============================================================
    # Hoja 3: Órdenes de Trabajo
    # ============================================================
    ws3 = wb.create_sheet('Ordenes_Trabajo')
    
    ordenes_df = dash_mod.get_ordenes_trabajo_data()
    for r in dataframe_to_rows(ordenes_df, index=False, header=True):
        ws3.append(r)
    
    apply_header_style(ws3, len(ordenes_df.columns))
    auto_width(ws3)
    
    # ============================================================
    # Hoja 4: Datos de Sensores
    # ============================================================
    ws4 = wb.create_sheet('Datos_Sensores')
    
    conn = get_connection()
    sensores_df = pd.read_sql_query('''
    SELECT d.*, e.codigo as equipo_codigo, e.nombre as equipo_nombre
    FROM datos_equipos d
    JOIN equipos e ON d.equipo_id = e.id
    ORDER BY d.fecha_hora DESC
    LIMIT 500
    ''', conn)
    conn.close()
    
    for r in dataframe_to_rows(sensores_df, index=False, header=True):
        ws4.append(r)
    
    apply_header_style(ws4, len(sensores_df.columns))
    auto_width(ws4)
    
    # ============================================================
    # Hoja 5: Repuestos
    # ============================================================
    ws5 = wb.create_sheet('Repuestos')
    
    conn = get_connection()
    repuestos_df = pd.read_sql_query('SELECT * FROM repuestos ORDER BY nombre', conn)
    conn.close()
    
    for r in dataframe_to_rows(repuestos_df, index=False, header=True):
        ws5.append(r)
    
    apply_header_style(ws5, len(repuestos_df.columns))
    auto_width(ws5)
    
    # ============================================================
    # Hoja 6: Gráficos Resumen
    # ============================================================
    ws6 = wb.create_sheet('Resumen_Graficos')
    
    # Datos para gráfico de estados
    ws6['A1'] = 'Estado'
    ws6['B1'] = 'Cantidad'
    for i, (estado, count) in enumerate(equipos_df['estado'].value_counts().items(), start=2):
        ws6[f'A{i}'] = estado
        ws6[f'B{i}'] = count
    
    # Gráfico de pie
    pie = PieChart()
    labels = Reference(ws6, min_col=1, min_row=2, max_row=1 + len(equipos_df['estado'].value_counts()))
    data = Reference(ws6, min_col=2, min_row=1, max_row=1 + len(equipos_df['estado'].value_counts()))
    pie.add_data(data, titles_from_data=True)
    pie.set_categories(labels)
    pie.title = 'Distribución de Estados de Equipos'
    ws6.add_chart(pie, 'D1')
    
    # Guardar en buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
